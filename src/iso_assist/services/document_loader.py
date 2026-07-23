"""
document_loader.py
──────────────────
Module 9.1 — Document Loader
Reads approved ISO documents (PDF, DOCX, TXT, Markdown) and returns
clean extracted text along with document metadata.

Supported formats: .pdf  .docx  .txt  .md  .markdown
"""

import json
import re
from pathlib import Path
from typing import Optional

from iso_assist.config import METADATA_PATH, SUPPORTED_EXTENSIONS


# ── Metadata helpers ──────────────────────────────────────────────────────────

def _load_metadata_registry() -> dict:
    """
    Load optional metadata registry from data/metadata/metadata.json.
    Returns empty dict if the file does not exist.

    Expected JSON format:
    {
      "Quality_Manual_v2.pdf": {
        "document_type": "Manual",
        "standard": "ISO 9001",
        "version": "2.0",
        "effective_date": "2024-01-01",
        "department": "Quality",
        "owner": "QA Manager",
        "status": "Approved"
      },
      ...
    }
    """
    if METADATA_PATH.exists():
        try:
            with open(METADATA_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def _build_metadata(file_path: Path, registry: dict) -> dict:
    """Combine file-level info with optional registry metadata."""
    base = {
        "document_name": file_path.name,
        "document_type": "Unknown",
        "standard": "Unknown",
        "version": "",
        "effective_date": "",
        "department": "",
        "owner": "",
        "status": "Approved",
    }
    # Overlay with registry entry if present
    entry = registry.get(file_path.name, {})
    base.update({k: v for k, v in entry.items() if v})
    return base


# ── Extractors ────────────────────────────────────────────────────────────────

def _extract_pdf(file_path: Path) -> str:
    """Extract text from a PDF using pdfplumber (preferred) or pypdf (fallback)."""
    text_pages = []
    try:
        import pdfplumber
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_pages.append(page_text)
        return "\n\n".join(text_pages)
    except Exception:
        pass

    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                text_pages.append(t)
        return "\n\n".join(text_pages)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from PDF '{file_path.name}': {e}")


def _extract_docx(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX '{file_path.name}': {e}")


def _extract_txt(file_path: Path) -> str:
    """Extract text from a plain text file."""
    try:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        raise RuntimeError(f"Failed to read TXT file '{file_path.name}': {e}")


def _extract_markdown(file_path: Path) -> str:
    """Extract plain text from a Markdown file (strip markdown syntax)."""
    try:
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        # Remove markdown formatting for cleaner embeddings
        raw = re.sub(r"#{1,6}\s*", "", raw)         # headings
        raw = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", raw)   # bold/italic
        raw = re.sub(r"`{1,3}[^`]*`{1,3}", "", raw)        # code
        raw = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", raw) # links
        return raw
    except Exception as e:
        raise RuntimeError(f"Failed to read Markdown file '{file_path.name}': {e}")


# ── Public API ────────────────────────────────────────────────────────────────

def load_document(file_path: Path, registry: Optional[dict] = None) -> dict:
    """
    Load a single document and return extracted text + metadata.

    Args:
        file_path: Path to the document file.
        registry:  Optional pre-loaded metadata registry dict.
                   If None, loads from disk automatically.

    Returns:
        {
            "text":     str,   # clean extracted text
            "metadata": dict,  # document_name, document_type, standard, etc.
        }

    Raises:
        ValueError:  If the file extension is not supported.
        RuntimeError: If text extraction fails.
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    # Extract text
    if ext == ".pdf":
        text = _extract_pdf(file_path)
    elif ext == ".docx":
        text = _extract_docx(file_path)
    elif ext == ".txt":
        text = _extract_txt(file_path)
    elif ext in (".md", ".markdown"):
        text = _extract_markdown(file_path)
    else:
        raise ValueError(f"Unhandled extension: {ext}")

    if not text.strip():
        raise RuntimeError(f"No text could be extracted from '{file_path.name}'. The file may be empty or image-only.")

    # Build metadata
    if registry is None:
        registry = _load_metadata_registry()
    metadata = _build_metadata(file_path, registry)

    return {"text": text, "metadata": metadata}


def load_all_documents(docs_path: Path) -> list[dict]:
    """
    Load all supported documents from a directory.

    Returns:
        List of dicts with keys 'text' and 'metadata'.
        Files that fail to load are skipped (error is logged to console).
    """
    docs_path = Path(docs_path)
    registry = _load_metadata_registry()
    results = []
    errors = []

    for file_path in sorted(docs_path.iterdir()):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            doc = load_document(file_path, registry=registry)
            results.append(doc)
            print(f"  [OK] Loaded: {file_path.name}")
        except Exception as e:
            errors.append((file_path.name, str(e)))
            print(f"  [SKIP] {file_path.name}: {e}")

    if errors:
        print(f"\nWarning: {len(errors)} file(s) could not be loaded.")

    return results
